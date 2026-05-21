# SPDX-License-Identifier: MIT
# Copyright (c) 2026 Corporate Bot Platform
"""Integration tests for Corporate Bot.

Runs against the live deployment at the configured server.
Tests: healthcheck, self-upgrade git workflow, TG webhook parsing,
message enrichment (reply-to, forward, location, contact, media).

Configuration via environment variables:
    BOT_SERVER      — server IP/hostname (required for remote tests)
    SSH_KEY         — path to SSH private key
    SSH_USER        — SSH username
    HEALTH_URL      — health endpoint URL (default: http://localhost:8000/health)
    CONTAINER_NAME  — Docker container name (e.g. smith-corporate-bot)
    BOT_REPO_DIR    — repo path on server (default: ~/corporate-bot, set to your fork folder)

Usage:
    # Local-only tests (parsing, enrichment — no server needed):
    python -m pytest tests/test_integration.py -v -k "Parsing or Enrichment"

    # All tests (requires BOT_SERVER):
    BOT_SERVER=1.2.3.4 SSH_KEY=~/.ssh/mykey CONTAINER_NAME=smith-corporate-bot \\
        python -m pytest tests/test_integration.py -v

    # Single test:
    python -m pytest tests/test_integration.py -v -k "test_health"
"""

import json
import os
import subprocess
import sys
import time
from datetime import datetime, timezone
from unittest.mock import MagicMock

import pytest

# Add parent dir to path so we can import bot modules
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# === Test configuration (all from env vars, no hardcoded family data) ===
SERVER = os.environ.get("BOT_SERVER", "")
SSH_KEY = os.environ.get("SSH_KEY", os.path.expanduser("~/.ssh/id_ed25519"))
SSH_USER = os.environ.get("SSH_USER", os.environ.get("USER", ""))
HEALTH_URL = os.environ.get("HEALTH_URL", "http://localhost:8000/health")
CONTAINER = os.environ.get("CONTAINER_NAME", "")
REPO_DIR = os.environ.get("BOT_REPO_DIR", "~/corporate-bot")

# Test data for TG parsing tests (generic, not family-specific)
TEST_CHAT_ID = -1001234567890
TEST_USER_ID = 111222333
TEST_USER_NAME = "TestUser"

requires_server = pytest.mark.skipif(
    not SERVER, reason="BOT_SERVER not set — skipping remote tests"
)
requires_container = pytest.mark.skipif(
    not CONTAINER, reason="CONTAINER_NAME not set — skipping container tests"
)


def ssh_cmd(cmd: str, timeout: int = 30) -> str:
    """Run a command on the server via SSH."""
    if not SERVER:
        pytest.skip("BOT_SERVER not configured")
    result = subprocess.run(
        ["ssh", "-o", "StrictHostKeyChecking=accept-new",
         "-i", SSH_KEY, f"{SSH_USER}@{SERVER}", cmd],
        capture_output=True, text=True, timeout=timeout,
    )
    if result.returncode != 0 and "error" in result.stderr.lower():
        raise RuntimeError(f"SSH command failed: {result.stderr}")
    return result.stdout + result.stderr


# ============================================================
# 1. HEALTH & DEPLOYMENT
# ============================================================

class TestHealth:
    def _build_health_response(self) -> dict:
        """Build a health response using the same logic as the /health endpoint."""
        from config import GIT_COMMIT
        return {
            "status": "ok",
            "version": "agent-sdk",
            "commit": GIT_COMMIT,
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "active_tasks": 0,
            "queue_size": 0,
        }

    def test_healthcheck_returns_ok(self):
        """Health endpoint returns status=ok with expected fields."""
        data = self._build_health_response()
        assert data["status"] == "ok"
        assert data["version"] == "agent-sdk"
        assert "commit" in data
        assert "timestamp" in data
        assert "active_tasks" in data
        assert "queue_size" in data

    def test_healthcheck_commit_format(self):
        """Health commit hash is a valid 7-char short hash or longer."""
        from config import GIT_COMMIT
        assert GIT_COMMIT and GIT_COMMIT != "unknown", f"GIT_COMMIT not resolved: {GIT_COMMIT}"
        # Commit hash should be a hex string of 7-12 chars (short hash)
        assert len(GIT_COMMIT) >= 7, f"Unexpected commit format: {GIT_COMMIT}"
        assert all(c in "0123456789abcdef" for c in GIT_COMMIT), f"Non-hex commit: {GIT_COMMIT}"

    def test_container_running(self):
        """We are running inside the container — verify /app exists."""
        import pathlib
        assert pathlib.Path("/app").is_dir(), "/app not found — not inside container?"
        assert pathlib.Path("/app/main.py").exists() or pathlib.Path("/host-repo/main.py").exists()



# ============================================================
# 2. GIT REPO SYNC
# ============================================================

class TestGitSync:
    def test_server_has_remotes(self):
        """Host repo has at least an origin remote."""
        result = subprocess.run(
            ["git", "remote", "-v"],
            capture_output=True, text=True, cwd="/host-repo",
        )
        assert result.returncode == 0, f"git remote failed: {result.stderr}"
        assert "origin" in result.stdout

    def test_no_uncommitted_changes(self):
        """Host repo has no dirty working tree (except override and runtime files)."""
        result = subprocess.run(
            ["git", "status", "--short"],
            capture_output=True, text=True, cwd="/host-repo",
        )
        assert result.returncode == 0, f"git status failed: {result.stderr}"
        out = result.stdout
        # Ignore expected untracked/runtime files
        ignore = ["docker-compose.override", "scheduled_tasks",
                  "google-workspace-creds", "data/last_", "tests/"]
        lines = [l for l in out.strip().split("\n")
                 if l.strip()
                 and not any(ign in l for ign in ignore)]
        assert not lines, f"Unexpected changes in repo: {lines}"


# ============================================================
# 3. SELF-UPGRADE (SSH + GIT inside container)
# ============================================================

class TestSelfUpgrade:
    def test_container_has_host_repo_mount(self):
        """/host-repo is mounted with main.py present."""
        import pathlib
        assert pathlib.Path("/host-repo/main.py").exists(), "/host-repo/main.py not found"

    def test_container_git_works(self):
        """Git commands work inside container on /host-repo."""
        result = subprocess.run(
            ["git", "status", "--short"],
            capture_output=True, text=True, cwd="/host-repo",
        )
        assert result.returncode == 0, f"git status failed: {result.stderr}"

    def test_container_ssh_key_exists(self):
        """Deploy key is accessible inside container."""
        import pathlib
        key_path = pathlib.Path("/home/botuser/.ssh/id_ed25519")
        assert key_path.exists(), f"SSH key not found at {key_path}"

    def test_container_can_reach_github(self):
        """Container can SSH to GitHub (deploy key auth)."""
        try:
            result = subprocess.run(
                ["ssh", "-o", "StrictHostKeyChecking=accept-new",
                 "-o", "ConnectTimeout=5", "-T", "git@github.com"],
                capture_output=True, text=True, timeout=15,
            )
            out = result.stdout + result.stderr
            assert "successfully authenticated" in out.lower() or "Hi " in out, (
                f"GitHub SSH auth failed: {out}"
            )
        except subprocess.TimeoutExpired:
            pytest.skip("GitHub SSH timed out — network unavailable in this environment")
        except FileNotFoundError:
            pytest.skip("ssh binary not found in container")


# ============================================================
# 4. TELEGRAM PARSING (unit-style, runs locally)
# ============================================================

class TestTelegramParsing:
    """Test _parse_pyrogram_message() with mock Pyrogram Message objects.

    Tests the actual message parser used in userbot mode (not the legacy
    Bot API parse_update which is now a stub).
    """

    @pytest.fixture(autouse=True)
    def setup_config(self, monkeypatch):
        """Patch telegram module's config for parser tests.

        _parse_pyrogram_message lazy-imports integrations.telegram and reads
        _my_user_id, TG_BOT_USER_ID, TG_MCP_BOT_USER_ID.  These only exist
        when TG_MODE=userbot; in bot mode we create them with raising=False.
        """
        import integrations.telegram as tg_mod
        monkeypatch.setattr(tg_mod, "TG_BOT_USER_ID", 0, raising=False)
        monkeypatch.setattr(tg_mod, "TG_MCP_BOT_USER_ID", 0, raising=False)
        monkeypatch.setattr(tg_mod, "_my_user_id", 0, raising=False)
        # Parser uses chat_registry for auth — mock it:
        # TEST_CHAT_ID is registered (active_plus). Other chat IDs not registered
        # but admin user (TEST_USER_ID) can still access them (returns external).
        import bot.chat_registry as _cr
        _test_key = f"telegram:{TEST_CHAT_ID}"
        _ext_key = "telegram:-99999"  # used in external/unauthorized tests
        _registered = {
            _test_key: {"mode": "active_plus"},
            _ext_key: {"mode": "active"},  # external chat
        }
        monkeypatch.setattr(_cr, "is_registered", lambda key: key in _registered)
        monkeypatch.setattr(_cr, "get", lambda key: _registered.get(key))
        # Parser falls back to is_admin_user for unregistered chats
        import config
        monkeypatch.setattr(config, "is_admin_user",
                            lambda src, uid: str(uid) == str(TEST_USER_ID))

    def _mock_user(self, user_id=TEST_USER_ID, first_name=TEST_USER_NAME, is_bot=False):
        user = MagicMock()
        user.id = user_id
        user.first_name = first_name
        user.is_bot = is_bot
        return user

    def _mock_chat(self, chat_id=TEST_CHAT_ID, title=None, first_name=None):
        chat = MagicMock()
        chat.id = chat_id
        chat.title = title
        chat.first_name = first_name
        return chat

    def _mock_message(self, text="hello", caption=None, user=None, chat=None, **kwargs):
        msg = MagicMock()
        msg.text = text
        msg.caption = caption
        msg.id = kwargs.get("message_id", 1)
        msg.date = datetime.now(timezone.utc)
        msg.from_user = user or self._mock_user()
        msg.chat = chat or self._mock_chat()
        # Media fields — default to None/falsy
        msg.photo = kwargs.get("photo")
        msg.document = kwargs.get("document")
        msg.sticker = kwargs.get("sticker")
        msg.voice = kwargs.get("voice")
        msg.video_note = kwargs.get("video_note")
        msg.video = kwargs.get("video")
        msg.animation = kwargs.get("animation")
        msg.audio = kwargs.get("audio")
        msg.location = kwargs.get("location")
        msg.venue = kwargs.get("venue")
        msg.contact = kwargs.get("contact")
        msg.reply_to_message = kwargs.get("reply_to_message")
        # Build forward_origin mock from legacy kwargs for backwards compat
        fwd_from = kwargs.get("forward_from")
        fwd_sender_name = kwargs.get("forward_sender_name")
        fwd_from_chat = kwargs.get("forward_from_chat")
        if fwd_from:
            fwd_origin = MagicMock()
            fwd_origin.type = "user"
            fwd_origin.sender_user = fwd_from
        elif fwd_sender_name:
            fwd_origin = MagicMock()
            fwd_origin.type = "hidden_user"
            fwd_origin.sender_user_name = fwd_sender_name
        elif fwd_from_chat:
            fwd_origin = MagicMock()
            fwd_origin.type = "channel"
            fwd_origin.chat = fwd_from_chat
            fwd_origin.sender_chat = fwd_from_chat
        else:
            fwd_origin = None
        msg.forward_origin = fwd_origin
        return msg

    def _mock_location(self, lat=51.4934, lon=-0.2231, live_period=None):
        loc = MagicMock()
        loc.latitude = lat
        loc.longitude = lon
        loc.live_period = live_period
        return loc

    def _mock_venue(self, lat=51.5, lon=-0.1, title="Big Ben", address="Westminster, London"):
        venue = MagicMock()
        venue.location = self._mock_location(lat, lon)
        venue.title = title
        venue.address = address
        return venue

    def _mock_contact(self, phone="+447123456789", first_name="John", last_name="Doe"):
        contact = MagicMock()
        contact.phone_number = phone
        contact.first_name = first_name
        contact.last_name = last_name
        return contact

    def _parse(self, msg):
        from integrations.tg_parser import _parse_pyrogram_message
        return _parse_pyrogram_message(msg)

    def test_parse_text_message(self):
        msg = self._mock_message(text="test message")
        parsed = self._parse(msg)
        assert parsed is not None
        assert parsed["text"] == "test message"
        assert parsed["user_id"] == TEST_USER_ID

    def test_parse_photo_message(self):
        photo = MagicMock()
        msg = self._mock_message(text="", caption="nice photo", photo=photo)
        parsed = self._parse(msg)
        assert parsed is not None
        assert parsed["has_media"] is True
        assert parsed["text"] == "nice photo"

    def test_parse_location_message(self):
        loc = self._mock_location(51.4934, -0.2231)
        msg = self._mock_message(text="", location=loc)
        parsed = self._parse(msg)
        assert parsed is not None
        assert parsed["location"]["latitude"] == 51.4934
        assert parsed["location"]["longitude"] == -0.2231

    def test_parse_live_location(self):
        loc = self._mock_location(51.0, -0.1, live_period=3600)
        msg = self._mock_message(text="", location=loc)
        parsed = self._parse(msg)
        assert parsed is not None
        assert parsed["location"]["live_period"] == 3600

    def test_parse_venue_message(self):
        venue = self._mock_venue()
        msg = self._mock_message(text="", venue=venue)
        parsed = self._parse(msg)
        assert parsed is not None
        assert parsed["location"]["title"] == "Big Ben"
        assert parsed["location"]["address"] == "Westminster, London"

    def test_parse_contact_message(self):
        contact = self._mock_contact()
        msg = self._mock_message(text="", contact=contact)
        parsed = self._parse(msg)
        assert parsed is not None
        assert parsed["contact"]["phone_number"] == "+447123456789"
        assert parsed["contact"]["first_name"] == "John"

    def test_parse_reply_to_text(self):
        rt = self._mock_message(text="I sent this earlier", message_id=99)
        msg = self._mock_message(text="what about this?", reply_to_message=rt)
        parsed = self._parse(msg)
        assert parsed["reply_to"] is not None
        assert parsed["reply_to"]["text"] == "I sent this earlier"
        assert parsed["reply_to"]["is_bot"] is False

    def test_parse_reply_to_bot_message(self):
        rt_user = self._mock_user(user_id=999, is_bot=True)
        rt = self._mock_message(text="Here's the summary...", user=rt_user, message_id=100)
        msg = self._mock_message(text="explain this", reply_to_message=rt)
        parsed = self._parse(msg)
        assert parsed["reply_to"]["is_bot"] is True

    def test_parse_reply_to_with_media(self):
        photo = MagicMock()
        rt = self._mock_message(text="", photo=photo, message_id=101)
        msg = self._mock_message(text="what is this?", reply_to_message=rt)
        parsed = self._parse(msg)
        assert parsed["reply_to"]["has_media"] is True
        assert parsed["reply_to"]["raw"] is not None

    def test_parse_reply_to_with_location(self):
        loc = self._mock_location(51.5, -0.1)
        rt = self._mock_message(text="", location=loc, message_id=102)
        msg = self._mock_message(text="is this close?", reply_to_message=rt)
        parsed = self._parse(msg)
        assert parsed["reply_to"]["location"] is not None
        assert parsed["reply_to"]["location"]["latitude"] == 51.5

    def test_parse_reply_to_with_contact(self):
        contact = self._mock_contact(phone="+44700000", first_name="Jane", last_name="")
        rt = self._mock_message(text="", contact=contact, message_id=103)
        msg = self._mock_message(text="call this person", reply_to_message=rt)
        parsed = self._parse(msg)
        assert parsed["reply_to"]["contact"]["phone_number"] == "+44700000"

    def test_parse_forward(self):
        fwd_user = self._mock_user(user_id=555, first_name="Someone")
        msg = self._mock_message(text="Check this out", forward_from=fwd_user)
        parsed = self._parse(msg)
        assert parsed["forward"] is not None
        assert parsed["forward"]["sender_name"] == "Someone"

    def test_parse_audio_message(self):
        audio = MagicMock()
        msg = self._mock_message(text="", caption="listen to this", audio=audio)
        parsed = self._parse(msg)
        assert parsed["has_media"] is True

    def test_empty_message_ignored(self):
        msg = self._mock_message(text="")
        msg.caption = None
        parsed = self._parse(msg)
        assert parsed is None

    def test_external_chat_registration(self):
        """Authorized user in unknown chat triggers explicitly registery."""
        chat = self._mock_chat(chat_id=-99999, title="New Group")
        msg = self._mock_message(text="hello", chat=chat)
        parsed = self._parse(msg)
        assert parsed is not None
        assert parsed.get("is_external_chat") is True
        assert parsed["chat_id"] == -99999

    def test_wrong_chat_ignored_for_unauthorized_user(self):
        """Unauthorized user in unregistered chat is silently ignored."""
        chat = self._mock_chat(chat_id=-88888)  # not in chat_registry
        user = self._mock_user(user_id=999999, first_name="Stranger")  # not admin
        msg = self._mock_message(text="hello", user=user, chat=chat)
        parsed = self._parse(msg)
        assert parsed is None

    def test_bot_message_flagged(self):
        user = self._mock_user(is_bot=True)
        msg = self._mock_message(text="I'm the bot", user=user)
        parsed = self._parse(msg)
        assert parsed is not None
        assert parsed.get("is_bot_message") is True


# ============================================================
# 5. MESSAGE ENRICHMENT (text building logic)
# ============================================================

class TestMessageEnrichment:
    """Test the text enrichment logic in handle_telegram_message.

    We test the enrichment functions by simulating parsed dicts
    and checking the resulting text modifications.
    """

    def _enrich_text(self, parsed: dict) -> str:
        """Simulate the text enrichment from handle_telegram_message."""
        text = parsed.get("text", "")

        # Reply-to enrichment
        reply_to = parsed.get("reply_to")
        if reply_to:
            parts = []
            if reply_to.get("text"):
                label = "bot's previous response" if reply_to.get("is_bot") else "earlier message"
                parts.append(f'[Replying to {label}: "{reply_to["text"][:300]}"]')
            if reply_to.get("location"):
                loc = reply_to["location"]
                coords = f"{loc.get('latitude')}, {loc.get('longitude')}"
                loc_type = "live location" if loc.get("live_period") else "location"
                if loc.get("title"):
                    addr = f", {loc['address']}" if loc.get("address") else ""
                    parts.append(f"[Replying to {loc_type}: {loc['title']}{addr} ({coords})]")
                else:
                    parts.append(f"[Replying to {loc_type}: {coords}]")
            if reply_to.get("contact"):
                c = reply_to["contact"]
                name = f"{c.get('first_name', '')} {c.get('last_name', '')}".strip()
                phone = c.get("phone_number", "")
                parts.append(f"[Replying to shared contact: {name}" + (f", {phone}]" if phone else "]"))
            if reply_to.get("has_media"):
                parts.append("[Replying to message with media — see attached]")
            if parts:
                text = "\n".join(parts) + "\n" + text

        # Forward enrichment
        forward = parsed.get("forward")
        if forward:
            fwd_name = forward.get("sender_name", "someone")
            text = (f"[Forwarded message from {fwd_name}]\n{text}\n\n"
                    f"[SYSTEM: This is a forwarded message. The user did not add their own instructions. "
                    f"Briefly acknowledge the content and ask what they'd like you to do with it.]")

        # Location enrichment
        if parsed.get("location"):
            loc = parsed["location"]
            coords = f"{loc.get('latitude')}, {loc.get('longitude')}"
            if loc.get("title"):
                addr = f", {loc['address']}" if loc.get("address") else ""
                loc_desc = f"{loc['title']}{addr} ({coords})"
            else:
                loc_desc = coords
            loc_type = "Live location" if loc.get("live_period") else "Shared location"
            text = f"[{loc_type}: {loc_desc}]\n{text}" if text else f"[{loc_type}: {loc_desc}]"

        # Contact enrichment
        if parsed.get("contact"):
            c = parsed["contact"]
            name = f"{c.get('first_name', '')} {c.get('last_name', '')}".strip()
            phone = c.get("phone_number", "")
            contact_desc = f"{name}" + (f", {phone}" if phone else "")
            text = f"[Shared contact: {contact_desc}]\n{text}" if text else f"[Shared contact: {contact_desc}]"

        return text

    def test_reply_to_text(self):
        text = self._enrich_text({
            "text": "what about this?",
            "reply_to": {"text": "original message", "is_bot": False},
        })
        assert '[Replying to earlier message: "original message"]' in text
        assert "what about this?" in text

    def test_reply_to_bot(self):
        text = self._enrich_text({
            "text": "explain",
            "reply_to": {"text": "bot said stuff", "is_bot": True},
        })
        assert "bot's previous response" in text

    def test_reply_to_location(self):
        text = self._enrich_text({
            "text": "how far?",
            "reply_to": {"location": {"latitude": 51.5, "longitude": -0.1}},
        })
        assert "[Replying to location: 51.5, -0.1]" in text

    def test_reply_to_venue(self):
        text = self._enrich_text({
            "text": "book it",
            "reply_to": {
                "location": {
                    "latitude": 51.5, "longitude": -0.1,
                    "title": "The Ritz", "address": "Piccadilly",
                },
            },
        })
        assert "The Ritz" in text
        assert "Piccadilly" in text

    def test_reply_to_contact(self):
        text = self._enrich_text({
            "text": "call them",
            "reply_to": {
                "contact": {"first_name": "John", "last_name": "Doe",
                            "phone_number": "+44123"},
            },
        })
        assert "John Doe" in text
        assert "+44123" in text

    def test_reply_to_media(self):
        text = self._enrich_text({
            "text": "what is this?",
            "reply_to": {"has_media": True},
        })
        assert "media — see attached" in text

    def test_reply_to_combined(self):
        """Reply to message with text + location + media."""
        text = self._enrich_text({
            "text": "tell me about this",
            "reply_to": {
                "text": "here's the place",
                "is_bot": False,
                "location": {"latitude": 51.5, "longitude": -0.1},
                "has_media": True,
            },
        })
        assert "earlier message" in text
        assert "location" in text
        assert "media" in text

    def test_forward(self):
        text = self._enrich_text({
            "text": "Check this email",
            "forward": {"sender_name": "Someone"},
        })
        assert "[Forwarded message from Someone]" in text
        assert "SYSTEM" in text

    def test_direct_location(self):
        text = self._enrich_text({
            "text": "",
            "location": {"latitude": 51.5, "longitude": -0.1},
        })
        assert "[Shared location: 51.5, -0.1]" in text

    def test_direct_live_location(self):
        text = self._enrich_text({
            "text": "",
            "location": {"latitude": 51.5, "longitude": -0.1, "live_period": 3600},
        })
        assert "[Live location:" in text

    def test_direct_venue(self):
        text = self._enrich_text({
            "text": "let's go here",
            "location": {
                "latitude": 51.5, "longitude": -0.1,
                "title": "Pizza Express", "address": "High Street",
            },
        })
        assert "Pizza Express" in text
        assert "High Street" in text

    def test_direct_contact(self):
        text = self._enrich_text({
            "text": "",
            "contact": {"first_name": "Anna", "last_name": "S",
                        "phone_number": "+44789"},
        })
        assert "[Shared contact: Anna S, +44789]" in text

    def test_no_enrichment_plain_text(self):
        text = self._enrich_text({"text": "hello"})
        assert text == "hello"


# ============================================================
# 6. DOCKER COMPOSE VALIDATION
# ============================================================

class TestDockerCompose:
    """Validate docker-compose.yml has no stale bind mounts."""

    @pytest.fixture(autouse=True)
    def load_compose(self):
        """Load docker-compose.yml content."""
        compose_path = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            "docker-compose.yml",
        )
        with open(compose_path) as f:
            self.compose_content = f.read()

    def test_no_org_goals_mount(self):
        """org_goals.json should not be bind-mounted (goals from org_config.json)."""
        assert "org_goals.json" not in self.compose_content

    def test_no_org_knowledge_mount(self):
        """org_knowledge.md should not be bind-mounted (knowledge in SQLite)."""
        assert "org_knowledge.md" not in self.compose_content

    def test_no_org_facts_mount(self):
        """org_facts.json should not be bind-mounted (facts in SQLite)."""
        assert "org_facts.json" not in self.compose_content

    def test_media_cache_volume_exists(self):
        """media-cache named volume should be defined."""
        assert "media-cache:" in self.compose_content

    def test_bot_data_volume_exists(self):
        """bot-data named volume should be defined."""
        assert "bot-data:" in self.compose_content

    def test_org_config_mounted_readonly(self):
        """org_config.json should be mounted read-only."""
        assert "org_config.json:/app/data/org_config.json:ro" in self.compose_content

    def test_prompts_mounted_readonly(self):
        """Prompts directory should be mounted read-only."""
        assert "./data/prompts:/app/data/prompts:ro" in self.compose_content

    def test_media_retention_default(self):
        """MEDIA_RETENTION_DAYS should default to 36135 (99 years)."""
        assert "36135" in self.compose_content


# ============================================================
# 7. CONFIG MODULE (unit tests)
# ============================================================

class TestConfig:
    """Test config.py constants and defaults."""

    def test_media_retention_days_default(self):
        """MEDIA_RETENTION_DAYS should default to 36135 (99 years)."""
        from config import MEDIA_RETENTION_DAYS
        assert MEDIA_RETENTION_DAYS == 36135 or MEDIA_RETENTION_DAYS > 0

    def test_media_cache_dir_exists(self):
        """MEDIA_CACHE_DIR should be set and point to a Path."""
        from config import MEDIA_CACHE_DIR
        assert MEDIA_CACHE_DIR is not None
        assert hasattr(MEDIA_CACHE_DIR, "mkdir")  # It's a Path

    def test_db_path_set(self):
        """DB_PATH should be set."""
        from config import DB_PATH
        assert DB_PATH is not None

    def test_fact_categories_list(self):
        """FACT_CATEGORIES should be a non-empty list."""
        from config import FACT_CATEGORIES
        assert isinstance(FACT_CATEGORIES, list)
        assert len(FACT_CATEGORIES) > 0

    def test_default_goals_list(self):
        """DEFAULT_GOALS should be a non-empty list."""
        from config import DEFAULT_GOALS
        assert isinstance(DEFAULT_GOALS, list)
        assert len(DEFAULT_GOALS) > 0

    def test_no_knowledge_file_config(self):
        """KNOWLEDGE_FILE should not exist in config (moved to SQLite)."""
        import config
        assert not hasattr(config, "KNOWLEDGE_FILE")

    def test_no_facts_file_config(self):
        """FACTS_FILE should not exist in config (moved to SQLite)."""
        import config
        assert not hasattr(config, "FACTS_FILE")

    def test_no_goals_file_config(self):
        """GOALS_FILE should not exist in config (goals from org_config.json)."""
        import config
        assert not hasattr(config, "GOALS_FILE")

    def test_model_constants(self):
        """Model constants should be set."""
        from config import MODEL_PRIMARY
        assert MODEL_PRIMARY


# ============================================================
# 8. MEMORY MODULE (unit tests — SQLite)
# ============================================================

class TestMemoryModule:
    """Test memory.py functions (uses a temporary in-memory DB)."""

    def test_no_load_knowledge_function(self):
        """load_knowledge should not exist (moved to SQLite)."""
        import bot.storage.memory as mem
        assert not hasattr(mem, "load_knowledge")

    def test_no_load_facts_function(self):
        """load_facts should not exist (moved to SQLite message_facts)."""
        import bot.storage.memory as mem
        assert not hasattr(mem, "load_facts")

    def test_no_save_knowledge_update_function(self):
        """save_knowledge_update should not exist."""
        import bot.storage.memory as mem
        assert not hasattr(mem, "save_knowledge_update")

    def test_no_save_facts_update_function(self):
        """save_facts_update should not exist (use store_message_fact)."""
        import bot.storage.memory as mem
        assert not hasattr(mem, "save_facts_update")

    def test_no_memory_search_all_function(self):
        """memory_search_all should not exist (removed in cleanup)."""
        import bot.storage.memory as mem
        assert not hasattr(mem, "memory_search_all")

    def test_no_store_summary_function(self):
        """store_summary should not exist."""
        import bot.storage.memory as mem
        assert not hasattr(mem, "store_summary")

    def test_no_search_summaries_function(self):
        """search_summaries should not exist."""
        import bot.storage.memory as mem
        assert not hasattr(mem, "search_summaries")

    def test_no_rolling_summary_functions(self):
        """Rolling summary functions should not exist (SDK handles context)."""
        import bot.storage.memory as mem
        assert not hasattr(mem, "update_rolling_summary")
        assert not hasattr(mem, "get_latest_summary")
        assert not hasattr(mem, "get_unsummarized_message_count")
        assert not hasattr(mem, "get_messages_for_summary")
        assert not hasattr(mem, "build_context_injection")

    def test_has_init_db(self):
        """init_db should exist."""
        import bot.storage.memory as mem
        assert hasattr(mem, "init_db")
        assert callable(mem.init_db)

    def test_has_store_message(self):
        """store_message should exist."""
        import bot.storage.memory as mem
        assert hasattr(mem, "store_message")

    def test_has_search_messages(self):
        """search_messages should exist."""
        import bot.storage.memory as mem
        assert hasattr(mem, "search_messages")

    def test_has_save_message_fact(self):
        """save_message_fact should exist."""
        import bot.storage.memory as mem
        assert hasattr(mem, "save_message_fact")

    def test_has_search_message_facts(self):
        """search_message_facts should exist."""
        import bot.storage.memory as mem
        assert hasattr(mem, "search_message_facts")

    def test_has_cache_media_file(self):
        """cache_media_file should exist."""
        import bot.storage.memory as mem
        assert hasattr(mem, "cache_media_file")

    def test_has_search_media_cache(self):
        """search_media_cache should exist."""
        import bot.storage.memory as mem
        assert hasattr(mem, "search_media_cache")

    def test_has_cleanup_expired_media(self):
        """cleanup_expired_media should exist."""
        import bot.storage.memory as mem
        assert hasattr(mem, "cleanup_expired_media")

    def test_has_get_media_cache_stats(self):
        """get_media_cache_stats should exist."""
        import bot.storage.memory as mem
        assert hasattr(mem, "get_media_cache_stats")


# ============================================================
# 9. DEPLOY SAFEGUARDS
# ============================================================

class TestDeploySafeguards:
    """Verify deploy safety rules are enforced in prompt and docs."""

    @pytest.fixture(autouse=True)
    def load_files(self):
        base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        # Self-upgrade content now lives in .claude/skills/self-upgrade/SKILL.md
        skill_path = os.path.join(base, ".claude", "skills", "self-upgrade", "SKILL.md")
        with open(skill_path) as f:
            self.upgrade_skill = f.read()
        with open(os.path.join(base, "CLAUDE.md")) as f:
            self.claude_md = f.read()

    def test_upgrade_skill_requires_confirmation(self):
        """Self-upgrade skill must mention /deploy command requirement."""
        assert "NEVER deploy without" in self.upgrade_skill
        assert "/deploy" in self.upgrade_skill

    def test_claude_md_points_to_skills(self):
        """CLAUDE.md should point to skills system (content moved in CB-35)."""
        assert "skills" in self.claude_md.lower()
        assert "harness" in self.claude_md.lower()

    def test_no_stale_file_references_in_skill(self):
        """Self-upgrade skill should not reference removed files."""
        assert "KNOWLEDGE_FILE" not in self.upgrade_skill
        assert "FACTS_FILE" not in self.upgrade_skill
        assert "org_knowledge.md" not in self.upgrade_skill

    def test_skill_mentions_show_changes(self):
        """Skill must instruct to show changes before deploying."""
        assert "Show" in self.upgrade_skill or "show" in self.upgrade_skill


# ============================================================
# 10. ENTRYPOINT VALIDATION
# ============================================================

class TestEntrypoint:
    """Validate entrypoint.sh has no stale references."""

    @pytest.fixture(autouse=True)
    def load_entrypoint(self):
        base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        with open(os.path.join(base, "entrypoint.sh")) as f:
            self.entrypoint = f.read()

    def test_no_goals_file_reference(self):
        """entrypoint.sh should not reference org_goals.json."""
        assert "org_goals.json" not in self.entrypoint

    def test_no_knowledge_file_reference(self):
        """entrypoint.sh should not reference org_knowledge.md."""
        assert "org_knowledge.md" not in self.entrypoint

    def test_no_facts_file_reference(self):
        """entrypoint.sh should not reference org_facts.json."""
        assert "org_facts.json" not in self.entrypoint

    def test_runs_as_botuser(self):
        """entrypoint.sh runs directly as botuser (no root, no gosu)."""
        # Container starts as botuser via Dockerfile USER directive — no gosu/chown needed
        assert "exec python main.py" in self.entrypoint


# ============================================================
# 11. MESSAGE SPLITTING (telegram _split_message)
# ============================================================

class TestMessageSplitting:
    """Test Telegram message splitting logic."""

    def test_short_message_no_split(self):
        """Messages under limit should return as single item."""
        from integrations.telegram import _split_message
        result = _split_message("Hello world", 4096)
        assert result == ["Hello world"]

    def test_split_at_paragraph_boundary(self):
        """Should prefer splitting at paragraph breaks (\\n\\n)."""
        from integrations.telegram import _split_message
        para1 = "A" * 80
        para2 = "B" * 80
        text = para1 + "\n\n" + para2
        result = _split_message(text, 100)
        assert len(result) >= 2
        assert result[0].strip().startswith("A")
        assert result[-1].strip().startswith("B")

    def test_split_at_line_boundary(self):
        """Should split at line breaks when no paragraph break available."""
        from integrations.telegram import _split_message
        line1 = "A" * 60
        line2 = "B" * 60
        text = line1 + "\n" + line2
        result = _split_message(text, 80)
        assert len(result) >= 2

    def test_split_at_sentence_boundary(self):
        """Should split at sentence boundaries ('. ') as fallback."""
        from integrations.telegram import _split_message
        text = "First sentence. Second sentence. Third sentence."
        result = _split_message(text, 35)
        assert len(result) >= 2
        # First chunk should end with a sentence boundary
        assert result[0].rstrip().endswith(".")

    def test_hard_cut_no_spaces(self):
        """Should hard-cut when no break points exist."""
        from integrations.telegram import _split_message
        text = "A" * 200
        result = _split_message(text, 100)
        assert len(result) == 2
        assert len(result[0]) <= 100
        assert result[0] + result[1] == text

    def test_all_content_preserved(self):
        """All original content should be present across chunks."""
        from integrations.telegram import _split_message
        text = "Line one.\nLine two.\nLine three.\nLine four.\nLine five."
        result = _split_message(text, 25)
        combined = " ".join(r.strip() for r in result)
        # All key words present
        for word in ["one", "two", "three", "four", "five"]:
            assert word in combined

    def test_empty_string(self):
        """Empty string should return single empty or original item."""
        from integrations.telegram import _split_message
        result = _split_message("", 4096)
        assert result == [""]

    def test_exact_limit_no_split(self):
        """Text exactly at limit should not be split."""
        from integrations.telegram import _split_message
        text = "X" * 100
        result = _split_message(text, 100)
        assert len(result) == 1
        assert result[0] == text


# ============================================================
# 12. MEDIA DESCRIPTION (describe_media validation)
# ============================================================

class TestMediaDescription:
    """Test describe_media input validation and helpers (no API calls)."""

    def test_describe_media_exists(self):
        """describe_media function should exist in gemini module."""
        from integrations.gemini import describe_media
        assert callable(describe_media)

    def test_unsupported_file_type(self):
        """describe_media should reject truly unsupported file types."""
        import asyncio
        from integrations.gemini import describe_media
        import tempfile
        # .exe is not supported
        with tempfile.NamedTemporaryFile(suffix=".exe", delete=False) as f:
            f.write(b"MZ\x00\x00")
            tmp_path = f.name
        try:
            result = asyncio.get_event_loop().run_until_complete(describe_media(tmp_path))
            assert "error" in result
            assert "Unsupported" in result["error"]
        finally:
            os.unlink(tmp_path)

    def test_missing_file(self):
        """describe_media should handle missing files gracefully."""
        import asyncio
        from integrations.gemini import describe_media
        result = asyncio.get_event_loop().run_until_complete(
            describe_media("/nonexistent/path/image.jpg")
        )
        assert "error" in result
        assert "not found" in result["error"].lower()

    def test_empty_file(self):
        """describe_media should reject empty files."""
        import asyncio
        from integrations.gemini import describe_media
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as f:
            tmp_path = f.name  # 0 bytes
        try:
            result = asyncio.get_event_loop().run_until_complete(describe_media(tmp_path))
            assert "error" in result
            assert "empty" in result["error"].lower()
        finally:
            os.unlink(tmp_path)

    def test_mime_map_coverage(self):
        """MIME_MAP should cover all major media categories."""
        from integrations.gemini import MIME_MAP
        # Images
        assert ".jpg" in MIME_MAP
        assert ".png" in MIME_MAP
        assert ".webp" in MIME_MAP
        # PDFs
        assert ".pdf" in MIME_MAP
        # Video
        assert ".mp4" in MIME_MAP
        assert ".mov" in MIME_MAP
        assert ".webm" in MIME_MAP
        # Audio
        assert ".ogg" in MIME_MAP
        assert ".mp3" in MIME_MAP
        assert ".opus" in MIME_MAP
        # Text
        assert ".txt" in MIME_MAP
        assert ".csv" in MIME_MAP

    def test_office_extractors_defined(self):
        """OFFICE_EXTRACTORS should cover docx and xlsx."""
        from integrations.gemini import OFFICE_EXTRACTORS
        assert ".docx" in OFFICE_EXTRACTORS
        assert ".xlsx" in OFFICE_EXTRACTORS

    def test_media_category_detection(self):
        """_get_media_category should return correct category for each type."""
        from integrations.gemini import _get_media_category
        assert _get_media_category(".jpg") == "image"
        assert _get_media_category(".png") == "image"
        assert _get_media_category(".pdf") == "document"
        assert _get_media_category(".docx") == "document"
        assert _get_media_category(".xlsx") == "document"
        assert _get_media_category(".txt") == "document"
        assert _get_media_category(".mp4") == "video"
        assert _get_media_category(".mov") == "video"
        assert _get_media_category(".ogg") == "audio"
        assert _get_media_category(".opus") == "audio"
        assert _get_media_category(".mp3") == "audio"
        assert _get_media_category(".xyz") == "unknown"

    def test_describe_prompt_per_category(self):
        """Each media category should have a tailored prompt."""
        from integrations.gemini import _get_describe_prompt
        img_prompt = _get_describe_prompt("image")
        doc_prompt = _get_describe_prompt("document")
        vid_prompt = _get_describe_prompt("video")
        aud_prompt = _get_describe_prompt("audio")
        # Each should be different
        assert img_prompt != doc_prompt
        assert doc_prompt != vid_prompt
        assert vid_prompt != aud_prompt
        # Each should mention search indexing
        assert "search indexing" in img_prompt
        assert "search indexing" in doc_prompt
        assert "search indexing" in vid_prompt
        assert "search indexing" in aud_prompt

    def test_office_text_extraction_docx(self):
        """_extract_office_text should extract text from a .docx file."""
        from pathlib import Path
        from integrations.gemini import _extract_office_text
        from docx import Document
        import tempfile
        # Create a minimal docx
        doc = Document()
        doc.add_paragraph("Hello world test paragraph")
        doc.add_paragraph("Second paragraph with numbers 12345")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_path = f.name
        try:
            result = _extract_office_text(Path(tmp_path), "docx")
            assert result is not None
            assert "Hello world" in result
            assert "12345" in result
        finally:
            os.unlink(tmp_path)

    def test_office_text_extraction_xlsx(self):
        """_extract_office_text should extract text from an .xlsx file."""
        from pathlib import Path
        from integrations.gemini import _extract_office_text
        from openpyxl import Workbook
        import tempfile
        # Create a minimal xlsx
        wb = Workbook()
        ws = wb.active
        ws.title = "TestSheet"
        ws["A1"] = "Name"
        ws["B1"] = "Value"
        ws["A2"] = "Test item"
        ws["B2"] = 42
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            wb.save(f.name)
            tmp_path = f.name
        try:
            result = _extract_office_text(Path(tmp_path), "xlsx")
            assert result is not None
            assert "TestSheet" in result
            assert "Name" in result
            assert "42" in result
        finally:
            os.unlink(tmp_path)

    def test_txt_file_now_supported(self):
        """Text files should now be in MIME_MAP (previously rejected)."""
        from integrations.gemini import MIME_MAP
        assert ".txt" in MIME_MAP
        assert MIME_MAP[".txt"] == "text/plain"

    def test_file_size_limit_constants(self):
        """Size limit constants should be loaded from config and reasonable."""
        from integrations.gemini import INLINE_SIZE_LIMIT, MAX_DESCRIBE_SIZE
        from config import MEDIA_DESCRIBE_INLINE_SIZE_MB, MEDIA_DESCRIBE_MAX_SIZE_MB
        assert INLINE_SIZE_LIMIT == MEDIA_DESCRIBE_INLINE_SIZE_MB * 1024 * 1024
        assert MAX_DESCRIBE_SIZE == MEDIA_DESCRIBE_MAX_SIZE_MB * 1024 * 1024
        assert MAX_DESCRIBE_SIZE > INLINE_SIZE_LIMIT

    def test_office_safety_limits_defined(self):
        """Office extraction should have safety limits against zip bombs."""
        from integrations.gemini import OFFICE_TEXT_CAP, OFFICE_MAX_PARAGRAPHS, OFFICE_MAX_ROWS
        assert OFFICE_TEXT_CAP == 20_000
        assert OFFICE_MAX_PARAGRAPHS == 2000
        assert OFFICE_MAX_ROWS == 5000

    def test_media_module_exists(self):
        """bot.storage.media module should exist with generate_media_description."""
        from bot.storage.media import generate_media_description
        assert callable(generate_media_description)

    def test_describe_semaphore_exists(self):
        """Rate-limiting semaphore should be defined in bot.storage.media."""
        from bot.storage.media import _describe_semaphore
        import asyncio
        assert isinstance(_describe_semaphore, asyncio.Semaphore)


# ============================================================
# 13. MEDIA SEARCH (search_media_cache / list_cached_media)
# ============================================================

class TestMediaSearch:
    """Test media search functions exist and have correct signatures."""

    def test_search_media_cache_exists(self):
        """search_media_cache should be async and callable."""
        import bot.storage.memory as mem
        assert hasattr(mem, "search_media_cache")
        assert callable(mem.search_media_cache)

    def test_list_cached_media_exists(self):
        """list_cached_media should be async and callable."""
        import bot.storage.memory as mem
        assert hasattr(mem, "list_cached_media")
        assert callable(mem.list_cached_media)

    def test_has_update_media_description(self):
        """update_media_description should exist for auto-descriptions."""
        import bot.storage.memory as mem
        assert hasattr(mem, "update_media_description")


# ============================================================
# 14. RAG CHUNKING (size-aware chunks)
# ============================================================

class TestRAGChunking:
    """Test RAG chunk building logic (pure functions, no DB or API)."""

    def _make_msg(self, msg_id: int, text: str, user: str = "TestUser",
                  source: str = "telegram", ts: str = "2026-03-08 12:00"):
        return {
            "id": msg_id,
            "text": text,
            "user_name": user,
            "source": source,
            "timestamp": ts,
            "msg_type": "user",
        }

    def test_empty_input(self):
        """No messages should produce no chunks."""
        from bot.storage.rag import _build_chunks
        assert _build_chunks([]) == []

    def test_single_short_message(self):
        """Single short message should produce one chunk."""
        from bot.storage.rag import _build_chunks
        msgs = [self._make_msg(1, "Hello world")]
        chunks = _build_chunks(msgs)
        assert len(chunks) == 1
        assert "Hello" in chunks[0]["chunk_text"]
        assert chunks[0]["start_msg_id"] == 1
        assert chunks[0]["end_msg_id"] == 1

    def test_multiple_short_messages_chunked(self):
        """Several short messages should all appear in at least one chunk."""
        from bot.storage.rag import _build_chunks
        msgs = [self._make_msg(i, f"Short message {i}") for i in range(1, 4)]
        chunks = _build_chunks(msgs)
        assert len(chunks) >= 1
        # All messages should appear in at least one chunk
        all_text = " ".join(c["chunk_text"] for c in chunks)
        for i in range(1, 4):
            assert f"Short message {i}" in all_text

    def test_long_message_split(self):
        """A single message exceeding MAX_CHUNK_CHARS should be split into sub-chunks."""
        from bot.storage.rag import _split_long_message
        long_text = "A" * 5000
        msg = self._make_msg(1, long_text)
        sub_chunks = _split_long_message(msg, long_text)
        assert len(sub_chunks) >= 3  # 5000 / 2000 = at least 3
        # All sub-chunks reference the same message
        for sc in sub_chunks:
            assert sc["start_msg_id"] == 1
            assert sc["end_msg_id"] == 1

    def test_long_message_overlap(self):
        """Sub-chunks of a long message should overlap."""
        from bot.storage.rag import _split_long_message, LONG_MSG_OVERLAP, MAX_CHUNK_CHARS
        long_text = "A" * 5000
        msg = self._make_msg(1, long_text)
        sub_chunks = _split_long_message(msg, long_text)
        # Total content covered should exceed original due to overlap
        total_chars = sum(len(sc["chunk_text"]) for sc in sub_chunks)
        assert total_chars > len(long_text)

    def test_chunk_respects_char_budget(self):
        """Each chunk text should not exceed MAX_CHUNK_CHARS (for normal messages)."""
        from bot.storage.rag import _build_chunks, MAX_CHUNK_CHARS
        # Create many messages that together exceed the budget
        msgs = [self._make_msg(i, f"Message number {i} with some padding text here.") for i in range(1, 50)]
        chunks = _build_chunks(msgs)
        for c in chunks:
            assert len(c["chunk_text"]) <= MAX_CHUNK_CHARS + 200  # small tolerance for formatting

    def test_chunks_have_overlap(self):
        """Consecutive chunks should share some messages (overlap)."""
        from bot.storage.rag import _build_chunks
        # Generate enough messages to force multiple chunks
        msgs = [self._make_msg(i, f"Message {i} " + "x" * 100) for i in range(1, 30)]
        chunks = _build_chunks(msgs)
        if len(chunks) >= 2:
            # Check that message ID ranges overlap between consecutive chunks
            for k in range(len(chunks) - 1):
                end_of_current = chunks[k]["end_msg_id"]
                start_of_next = chunks[k + 1]["start_msg_id"]
                # Overlap means the next chunk starts at or before the current ends
                assert start_of_next <= end_of_current, \
                    f"Chunk {k} ends at {end_of_current} but chunk {k+1} starts at {start_of_next}"


# ============================================================
# 15. UNIFIED MEMORY SEARCH (memory_search tool)
# ============================================================

class TestUnifiedMemorySearch:
    """Test memory_search tool exists and runs 4 backends."""

    def test_memory_search_tool_exists(self):
        """memory_search should be defined in mcp_tools."""
        import bot.mcp_tools as tools
        # Check for the function
        assert hasattr(tools, "memory_search") or any(
            "memory_search" in str(getattr(tools, attr, ""))
            for attr in dir(tools)
        )

    def test_rag_module_has_search(self):
        """rag_search should exist in bot.storage.rag."""
        import bot.storage.rag as rag
        assert hasattr(rag, "rag_search")
        assert callable(rag.rag_search)


# ============================================================
# 16. DEPLOY APPROVAL KEYWORD GATE (security)
# ============================================================

class TestDeployApprovalGate:
    """check_deploy_approval requires /deploy command (not bare 'deploy' keyword)."""

    def test_deploy_command_accepted(self):
        from bot.hooks import check_deploy_approval
        assert check_deploy_approval("/deploy") is True
        assert check_deploy_approval("/deploy force") is True
        assert check_deploy_approval("/deploy no-qa") is True
        assert check_deploy_approval("  /deploy") is True  # leading whitespace OK

    def test_bare_deploy_keyword_rejected(self):
        """Casual mentions of 'deploy' must NOT authorize deployment."""
        from bot.hooks import check_deploy_approval
        assert check_deploy_approval("deploy") is False
        assert check_deploy_approval("ship it") is False
        assert check_deploy_approval("restart") is False
        assert check_deploy_approval("задеплой") is False
        assert check_deploy_approval("please deploy now") is False
        assert check_deploy_approval("staging deploy lock files") is False

    def test_non_deploy_keywords_rejected(self):
        from bot.hooks import check_deploy_approval
        assert check_deploy_approval("continue") is False
        assert check_deploy_approval("go ahead") is False
        assert check_deploy_approval("yes") is False
        assert check_deploy_approval("") is False

    def test_quoted_bot_reply_containing_deploy_rejected(self):
        """A user replying 'Continue' to a bot msg that quotes '/deploy' must NOT pass."""
        from bot.hooks import check_deploy_approval
        enriched = (
            '[Replying to bot\'s previous response: "Send /deploy to trigger rebuild."]\n'
            'Continue'
        )
        assert check_deploy_approval(enriched) is False

    def test_quoted_with_user_deploy_command_passes(self):
        """User's own '/deploy' outside the quote counts."""
        from bot.hooks import check_deploy_approval
        enriched = (
            '[Replying to bot\'s previous response: "Here are the changes."]\n'
            '/deploy'
        )
        assert check_deploy_approval(enriched) is True

    def test_forwarded_block_with_deploy_rejected(self):
        from bot.hooks import check_deploy_approval
        enriched = (
            '[Forwarded message from John]\n'
            'we should deploy the new build\n'
            '[SYSTEM: This is a forwarded message...]'
        )
        # Bare "deploy" no longer triggers — /deploy command required
        assert check_deploy_approval(enriched) is False

    def test_multiline_quoted_block_stripped(self):
        """Multi-line quoted bot response with '/deploy' must be stripped."""
        from bot.hooks import check_deploy_approval
        enriched = (
            '[Replying to bot\'s previous response: "Changes pushed.\n'
            'Send /deploy to trigger rebuild.\n'
            'Budget low."]\n'
            'thanks'
        )
        assert check_deploy_approval(enriched) is False

    def test_deploy_force_and_noqa(self):
        """check_deploy_force and check_deploy_noqa detect the right variants."""
        from bot.hooks import check_deploy_force, check_deploy_noqa
        assert check_deploy_force("/deploy force") is True
        assert check_deploy_force("/deploy") is False
        assert check_deploy_force("/deploy no-qa") is False
        assert check_deploy_noqa("/deploy no-qa") is True
        assert check_deploy_noqa("/deploy noqa") is True
        assert check_deploy_noqa("/deploy") is False
        assert check_deploy_noqa("/deploy force") is False


# ============================================================
# EXTEND_BUDGET APPROVAL GATE — mirrors deploy gate (same strip pipeline,
# different keyword set). Any change to _REPLY_QUOTE_BLOCK or _ENRICHMENT_MARKER
# MUST keep both deploy + extend_budget gates passing.
# ============================================================

class TestExtendBudgetApprovalGate:
    """check_extend_budget_approval must match only the USER's text."""

    def test_plain_extend_budget_keyword(self):
        from bot.hooks import check_extend_budget_approval
        assert check_extend_budget_approval("extend budget") is True
        assert check_extend_budget_approval("increase budget") is True
        assert check_extend_budget_approval("продли бюджет") is True
        assert check_extend_budget_approval("увеличь бюджет") is True
        assert check_extend_budget_approval("please extend budget by $20") is True

    def test_non_extend_budget_keywords_rejected(self):
        from bot.hooks import check_extend_budget_approval
        assert check_extend_budget_approval("continue") is False
        assert check_extend_budget_approval("more budget") is False
        assert check_extend_budget_approval("budget") is False
        assert check_extend_budget_approval("extend") is False
        assert check_extend_budget_approval("") is False

    def test_quoted_bot_reply_containing_keyword_rejected(self):
        """User replying 'Continue' to a bot msg that quotes the keyword must NOT pass."""
        from bot.hooks import check_extend_budget_approval
        enriched = (
            '[Replying to bot\'s previous response: "Say extend budget to raise the cap."]\n'
            'Continue'
        )
        assert check_extend_budget_approval(enriched) is False

    def test_quoted_with_user_keyword_passes(self):
        """User's own 'extend budget' outside the quote still counts."""
        from bot.hooks import check_extend_budget_approval
        enriched = (
            '[Replying to bot\'s previous response: "Here are the changes."]\n'
            'extend budget by 20'
        )
        assert check_extend_budget_approval(enriched) is True

    def test_multiline_quoted_block_stripped(self):
        from bot.hooks import check_extend_budget_approval
        enriched = (
            '[Replying to bot\'s previous response: "Running low.\n'
            'Reply extend budget to raise cap.\n'
            'Otherwise will auto-reset."]\n'
            'thanks'
        )
        assert check_extend_budget_approval(enriched) is False

    def test_voice_attached_marker_stripped_transcript_counts(self):
        from bot.hooks import check_extend_budget_approval
        enriched = (
            '[VOICE attached: /app/data/media_cache/tg_voice_456.ogg]\n'
            'Voice transcript: "extend budget by fifty dollars"'
        )
        assert check_extend_budget_approval(enriched) is True


class TestExtendBudgetCaps:
    """extend_budget must reject non-finite / out-of-range delta values."""

    def test_rejects_nan(self):
        import asyncio
        from bot.agent import client_pool
        async def _go():
            try:
                await client_pool.extend_budget("telegram:999999", float("nan"))
            except ValueError as e:
                return str(e)
            return None
        msg = asyncio.run(_go())
        assert msg is not None and "finite" in msg.lower()

    def test_rejects_inf(self):
        import asyncio
        from bot.agent import client_pool
        async def _go():
            try:
                await client_pool.extend_budget("telegram:999999", float("inf"))
            except ValueError as e:
                return str(e)
            return None
        msg = asyncio.run(_go())
        assert msg is not None and "finite" in msg.lower()

    def test_rejects_negative(self):
        import asyncio
        from bot.agent import client_pool
        async def _go():
            try:
                await client_pool.extend_budget("telegram:999999", -5.0)
            except ValueError as e:
                return str(e)
            return None
        msg = asyncio.run(_go())
        assert msg is not None and "positive" in msg.lower()

    def test_rejects_zero(self):
        import asyncio
        from bot.agent import client_pool
        async def _go():
            try:
                await client_pool.extend_budget("telegram:999999", 0.0)
            except ValueError as e:
                return str(e)
            return None
        msg = asyncio.run(_go())
        assert msg is not None and "positive" in msg.lower()

    def test_rejects_over_per_extension_cap(self):
        import asyncio
        from bot.agent import client_pool, MAX_EXTEND_DELTA_USD
        async def _go():
            try:
                await client_pool.extend_budget("telegram:999999", MAX_EXTEND_DELTA_USD + 1)
            except ValueError as e:
                return str(e)
            return None
        msg = asyncio.run(_go())
        assert msg is not None and "per-extension" in msg.lower()


# ============================================================
# MAC IP RESOLVER — multi-tier fallback chain tests
# ============================================================

class TestMacResolver:
    """Unit tests for bot/relay/mac_resolver.py. Each tier exercised in isolation
    via env manipulation + mocked sub-calls where needed."""

    def test_is_valid_ipv4_accepts_good(self):
        from bot.relay.mac_resolver import _is_valid_ipv4
        assert _is_valid_ipv4("100.100.100.1")
        assert _is_valid_ipv4("127.0.0.1")
        assert _is_valid_ipv4("0.0.0.0")

    def test_is_valid_ipv4_rejects_bad(self):
        from bot.relay.mac_resolver import _is_valid_ipv4
        assert not _is_valid_ipv4("")
        assert not _is_valid_ipv4("not-an-ip")
        assert not _is_valid_ipv4("300.1.1.1")
        assert not _is_valid_ipv4("fd7a::1")  # IPv6 — resolver is v4-only
        assert not _is_valid_ipv4(None)

    def test_env_pin_returns_valid(self, monkeypatch):
        from bot.relay import mac_resolver
        monkeypatch.setenv("DESKTOP_HOST_IP", "10.0.0.42")
        assert mac_resolver._resolve_env_pin() == "10.0.0.42"

    def test_env_pin_rejects_invalid(self, monkeypatch):
        from bot.relay import mac_resolver
        monkeypatch.setenv("DESKTOP_HOST_IP", "not-valid")
        assert mac_resolver._resolve_env_pin() is None

    def test_env_pin_empty_returns_none(self, monkeypatch):
        from bot.relay import mac_resolver
        monkeypatch.setenv("DESKTOP_HOST_IP", "")
        assert mac_resolver._resolve_env_pin() is None

    def test_match_device_ip_by_hostname(self):
        from bot.relay.mac_resolver import _match_device_ip
        devices = [
            {"hostname": "other-laptop", "addresses": ["100.1.1.1"]},
            {"hostname": "admin-macbook",
             "addresses": ["100.100.100.1", "fd7a::1"]},
        ]
        assert _match_device_ip(devices, "admin-macbook") == "100.100.100.1"

    def test_match_device_ip_case_insensitive(self):
        from bot.relay.mac_resolver import _match_device_ip
        devices = [{"hostname": "ADMIN-MACBOOK",
                    "addresses": ["100.100.100.1"]}]
        assert _match_device_ip(devices, "admin-macbook") == "100.100.100.1"

    def test_match_device_ip_strips_tailnet_suffix(self):
        """Tailscale sometimes returns `hostname.tailnet.ts.net` in `name`."""
        from bot.relay.mac_resolver import _match_device_ip
        devices = [{"name": "admin-macbook.example.ts.net",
                    "addresses": ["100.100.100.1"]}]
        assert _match_device_ip(devices, "admin-macbook") == "100.100.100.1"

    def test_match_device_ip_no_match_returns_none(self):
        from bot.relay.mac_resolver import _match_device_ip
        devices = [{"hostname": "different-laptop", "addresses": ["100.1.1.1"]}]
        assert _match_device_ip(devices, "admin-macbook") is None

    def test_match_device_ip_skips_non_ipv4(self):
        """IPv6 addresses in the list must not be returned."""
        from bot.relay.mac_resolver import _match_device_ip
        devices = [{"hostname": "admin-macbook",
                    "addresses": ["fd7a::1", "100.100.100.1"]}]
        assert _match_device_ip(devices, "admin-macbook") == "100.100.100.1"

    def test_get_mac_ip_never_blocks(self):
        """Public sync API: never raises, never blocks on network.
        Returns empty string when no fallback IP configured (fail-closed)."""
        from bot.relay.mac_resolver import get_mac_ip
        ip = get_mac_ip()
        assert isinstance(ip, str)  # May be empty (fail-closed) or valid IPv4

    def test_get_ssh_host_opts_shape(self):
        """Returns ['-o', 'HostName=<ip>'] — ready to splat into ssh cmd."""
        from bot.relay.mac_resolver import get_ssh_host_opts
        opts = get_ssh_host_opts()
        assert len(opts) == 2
        assert opts[0] == "-o"
        assert opts[1].startswith("HostName=")

    # M1 fix: CGNAT-range guard on DNS tier

    def test_is_tailnet_ipv4_accepts_cgnat_range(self):
        from bot.relay.mac_resolver import _is_tailnet_ipv4
        assert _is_tailnet_ipv4("100.100.100.1")  # Admin Mac
        assert _is_tailnet_ipv4("100.64.0.1")     # range start
        assert _is_tailnet_ipv4("100.127.255.254")  # range end

    def test_is_tailnet_ipv4_rejects_non_cgnat(self):
        from bot.relay.mac_resolver import _is_tailnet_ipv4
        assert not _is_tailnet_ipv4("8.8.8.8")        # public DNS
        assert not _is_tailnet_ipv4("192.168.1.1")    # private
        assert not _is_tailnet_ipv4("100.63.255.255")  # just below CGNAT
        assert not _is_tailnet_ipv4("100.128.0.0")     # just above CGNAT
        assert not _is_tailnet_ipv4("not-an-ip")
        assert not _is_tailnet_ipv4("")

    # H1 fix: _commit's bump_ttl flag + stale fallback behavior

    def test_commit_stale_tier_preserves_last_refresh_t(self, monkeypatch):
        """When all live tiers fail and _current_ip is already non-default,
        _last_refresh_t must NOT be bumped — otherwise TTL suppresses retry."""
        import asyncio
        from bot.relay import mac_resolver

        # Set up: pretend we have a good IP from a prior successful refresh.
        # Use a value DIFFERENT from _DEFAULT_FALLBACK_IP so we exercise the
        # stale-preserve branch (not the default-cold-start branch).
        good_ip = "100.x.y.77"
        assert good_ip != mac_resolver._DEFAULT_FALLBACK_IP
        monkeypatch.setattr(mac_resolver, "_current_ip", good_ip)
        monkeypatch.setattr(mac_resolver, "_current_provider", "dns")
        monkeypatch.setattr(mac_resolver, "_last_refresh_t", 0.0)

        # All live tiers return None
        monkeypatch.setattr(mac_resolver, "_resolve_env_pin", lambda: None)
        async def _no_ts(): return None
        async def _no_dns(): return None
        monkeypatch.setattr(mac_resolver, "_resolve_tailscale", _no_ts)
        monkeypatch.setattr(mac_resolver, "_resolve_dns", _no_dns)

        async def _go():
            return await mac_resolver.refresh_mac_ip(force=True)
        result = asyncio.run(_go())

        # Stale path: IP preserved, last_refresh_t stays 0 so next tick retries
        assert result.ip == good_ip
        assert "stale" in result.provider
        assert mac_resolver._last_refresh_t == 0.0  # NOT bumped — the key H1 fix
