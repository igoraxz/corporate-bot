# SPDX-License-Identifier: MIT
# Copyright (c) 2026 Corporate Bot Platform
"""Gemini API integration — image generation (Imagen) + image editing (Gemini native) + media descriptions."""

import asyncio
import logging
import time
from pathlib import Path

from config import GEMINI_API_KEY, TMP_DIR, VOICE_TRANSCRIBE_TIMEOUT_S, VOICE_TRANSCRIBE_MAX_CHARS

log = logging.getLogger(__name__)

# Gemini models — configurable via env vars for easy swaps on deprecation
import os
MODEL_FLASH = os.environ.get("GEMINI_MODEL_FLASH", "gemini-3.1-flash-image-preview")  # Image gen/edit (fast)
# MODEL_PRO defaults to same as FLASH — no stable Pro image model available as of Mar 2026.
# Override via GEMINI_MODEL_PRO env var when a suitable Pro model launches.
MODEL_PRO = os.environ.get("GEMINI_MODEL_PRO", "gemini-3.1-flash-image-preview")
# Text-only model for media descriptions (cheap, fast)
MODEL_FLASH_TEXT = os.environ.get("GEMINI_MODEL_DESCRIBE", "gemini-2.5-flash-lite")


async def generate_image(prompt: str, filename: str = "generated.png") -> dict:
    """Generate an image using Google Imagen 4.0.

    Returns dict with 'file_path' on success or 'error' on failure.
    """
    if not GEMINI_API_KEY:
        return {"error": "GEMINI_API_KEY not configured"}

    try:
        from google import genai
        from google.genai import types
    except ImportError:
        return {"error": "google-genai package not installed. Run: pip install google-genai"}

    client = genai.Client(api_key=GEMINI_API_KEY)

    try:
        response = await asyncio.get_running_loop().run_in_executor(
            None,
            lambda: client.models.generate_images(
                model="imagen-4.0-generate-001",
                prompt=prompt,
                config=types.GenerateImagesConfig(
                    number_of_images=1,
                    aspect_ratio="1:1",
                    person_generation="DONT_ALLOW",
                ),
            ),
        )
    except Exception as e:
        log.error(f"Imagen API error: {e}")
        return {"error": f"Imagen API error: {str(e)[:300]}"}

    if not response.generated_images:
        return {"error": "Imagen returned no images"}

    image = response.generated_images[0].image

    # Sanitize filename — strip path components to prevent traversal (../../.env)
    safe_name = Path(filename).name
    if not safe_name or safe_name in (".", ".."):
        safe_name = "generated.png"
    file_path = TMP_DIR / safe_name
    file_path.write_bytes(image.image_bytes)

    log.info(f"Generated image: {file_path} ({file_path.stat().st_size / 1024:.0f}KB)")
    return {"file_path": str(file_path)}


async def edit_image(
    image_path: str,
    prompt: str,
    filename: str = "edited.png",
    use_pro: bool = False,
) -> dict:
    """Edit an existing image using Gemini's native image generation.

    Sends the image + text prompt to Gemini, which returns a modified image.
    Uses flash model by default (fast), pro model for complex edits.

    Returns dict with 'file_path' on success or 'error' on failure.
    """
    if not GEMINI_API_KEY:
        return {"error": "GEMINI_API_KEY not configured"}

    try:
        from google import genai
        from google.genai import types
        from PIL import Image
    except ImportError as e:
        return {"error": f"Missing package: {e}. Need: google-genai, Pillow"}

    src = Path(image_path)
    if not src.exists():
        return {"error": f"Source image not found: {image_path}"}

    client = genai.Client(api_key=GEMINI_API_KEY)
    model = MODEL_PRO if use_pro else MODEL_FLASH

    try:
        img = Image.open(src)
        response = await asyncio.get_running_loop().run_in_executor(
            None,
            lambda: client.models.generate_content(
                model=model,
                contents=[prompt, img],
                config=types.GenerateContentConfig(
                    response_modalities=["TEXT", "IMAGE"],
                ),
            ),
        )
    except Exception as e:
        log.error(f"Gemini image edit error: {e}")
        return {"error": f"Gemini image edit error: {str(e)[:300]}"}

    # Extract generated image from response
    if not response.candidates or not response.candidates[0].content.parts:
        return {"error": "Gemini returned no image output"}

    for part in response.candidates[0].content.parts:
        if part.inline_data is not None:
            # Sanitize filename — strip path components to prevent traversal
            safe_name = Path(filename).name
            if not safe_name or safe_name in (".", ".."):
                safe_name = "edited.png"
            file_path = TMP_DIR / safe_name
            file_path.write_bytes(part.inline_data.data)
            log.info(f"Edited image: {file_path} ({file_path.stat().st_size / 1024:.0f}KB)")
            return {"file_path": str(file_path)}

    return {"error": "Gemini response contained no image data"}


# --- Media description prompts (per media category) ---

DESCRIBE_PROMPT_IMAGE = (
    "Describe this image concisely for search indexing. Include: "
    "1) What the image shows (objects, people, scene, text visible) "
    "2) Any text/numbers visible (OCR) "
    "3) Document type if applicable (receipt, letter, screenshot, etc.) "
    "Keep the description under 500 characters. Be factual, not creative."
)

DESCRIBE_PROMPT_DOCUMENT = (
    "Summarize this document concisely for search indexing. Include: "
    "1) Document type (letter, invoice, report, form, etc.) "
    "2) Key topics, names, dates, and numbers mentioned "
    "3) Any actionable items or deadlines "
    "Keep the summary under 500 characters. Be factual, not creative."
)

DESCRIBE_PROMPT_VIDEO = (
    "Describe this video concisely for search indexing. Include: "
    "1) What happens in the video (actions, scenes, setting) "
    "2) Any speech or dialogue (brief transcript of key points) "
    "3) Any text visible on screen "
    "Keep the description under 500 characters. Be factual, not creative."
)

DESCRIBE_PROMPT_AUDIO = (
    "Transcribe and describe this audio for search indexing. Include: "
    "1) Full transcript of speech (if short) or key points (if long) "
    "2) Language spoken "
    "3) Type of audio (voice message, music, conversation, etc.) "
    "Keep the description under 800 characters. Be factual, not creative."
)

DESCRIBE_PROMPT_PDF = (
    "Extract ALL text content from this PDF document. Rules:\n"
    "1. Preserve the structure: headings, bullet points, numbered lists\n"
    "2. Include ALL text from every page/slide — do not skip or summarize\n"
    "3. For presentations: label each slide (Slide 1, Slide 2, etc.)\n"
    "4. Include any visible numbers, dates, names, URLs\n"
    "5. For charts/graphs: describe type and key data points\n"
    "6. For images: briefly note what they show\n"
    "7. Output as clean readable text\n"
    "Be thorough — extract everything visible. Do not summarize."
)

# --- MIME type mapping for Gemini-native media ---

MIME_MAP: dict[str, str] = {
    # Images (Gemini native)
    ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png",
    ".webp": "image/webp", ".bmp": "image/bmp", ".heic": "image/heic",
    ".heif": "image/heif",
    # PDF (Gemini native)
    ".pdf": "application/pdf",
    # Video (Gemini native)
    ".mp4": "video/mp4", ".mov": "video/quicktime", ".avi": "video/avi",
    ".webm": "video/webm", ".mpeg": "video/mpeg", ".3gp": "video/3gpp",
    # Audio (Gemini native)
    ".ogg": "audio/ogg", ".mp3": "audio/mp3", ".wav": "audio/wav",
    ".m4a": "audio/m4a", ".aac": "audio/aac", ".flac": "audio/flac",
    ".opus": "audio/opus",
    # Text (Gemini native)
    ".txt": "text/plain", ".csv": "text/csv", ".html": "text/html",
    ".json": "application/json", ".xml": "text/xml",
}

# Office formats handled via text extraction (not Gemini-native)
OFFICE_EXTRACTORS: dict[str, str] = {
    ".docx": "docx",
    ".xlsx": "xlsx",
}

# Size thresholds loaded from config (externalized for tuning without code changes)
# MEDIA_DESCRIBE_INLINE_SIZE_MB: files under this use Part.from_bytes (fast, no upload)
# MEDIA_DESCRIBE_MAX_SIZE_MB: files over this are skipped (cost/time protection)
# Defaults: 20 MB inline threshold, 100 MB max (covers all team files/docs/voice,
# excludes feature-length videos that would cost ~$0.15+ and take minutes to process)
from config import (
    MEDIA_DESCRIBE_MAX_SIZE_MB, MEDIA_DESCRIBE_INLINE_SIZE_MB, MEDIA_DESCRIBE_TIMEOUT_S,
)
INLINE_SIZE_LIMIT = MEDIA_DESCRIBE_INLINE_SIZE_MB * 1024 * 1024
MAX_DESCRIBE_SIZE = MEDIA_DESCRIBE_MAX_SIZE_MB * 1024 * 1024

# Max chars to extract from office documents (20k chars ≈ 8 pages, ~5k tokens, <$0.001)
OFFICE_TEXT_CAP = 20_000
# Safety limits for office doc extraction (prevent zip bombs / huge files)
OFFICE_MAX_PARAGRAPHS = 2000  # max paragraphs from .docx
OFFICE_MAX_ROWS = 5000  # max rows from .xlsx


def _get_media_category(suffix: str) -> str:
    """Determine media category from file extension."""
    if suffix in {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".heic", ".heif"}:
        return "image"
    if suffix in {".pdf", ".txt", ".csv", ".html", ".json", ".xml", ".docx", ".xlsx"}:
        return "document"
    if suffix in {".mp4", ".mov", ".avi", ".webm", ".mpeg", ".3gp"}:
        return "video"
    if suffix in {".ogg", ".mp3", ".wav", ".m4a", ".aac", ".flac", ".opus"}:
        return "audio"
    return "unknown"


def _get_describe_prompt(category: str, suffix: str = "") -> str:
    """Return the appropriate description prompt for a media category.

    PDFs get a dedicated full-extraction prompt instead of the short document summary.
    """
    if suffix == ".pdf":
        return DESCRIBE_PROMPT_PDF
    return {
        "image": DESCRIBE_PROMPT_IMAGE,
        "document": DESCRIBE_PROMPT_DOCUMENT,
        "video": DESCRIBE_PROMPT_VIDEO,
        "audio": DESCRIBE_PROMPT_AUDIO,
    }.get(category, DESCRIBE_PROMPT_IMAGE)


def _extract_office_text(file_path: Path, fmt: str) -> str | None:
    """Extract text content from office documents. Returns text or None.

    Safety limits prevent zip bombs and huge files from consuming resources:
    - Paragraph/row count caps stop iteration before memory exhaustion
    - Character cap limits the final output size
    """
    try:
        if fmt == "docx":
            from docx import Document
            doc = Document(str(file_path))
            paragraphs = []
            count = 0
            for p in doc.paragraphs:
                if p.text.strip():
                    paragraphs.append(p.text)
                    count += 1
                    if count >= OFFICE_MAX_PARAGRAPHS:
                        paragraphs.append("[... truncated]")
                        break
            # Also extract table content (within same limit)
            for table in doc.tables:
                if count >= OFFICE_MAX_PARAGRAPHS:
                    break
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append("\t".join(cells))
                        count += 1
                        if count >= OFFICE_MAX_PARAGRAPHS:
                            paragraphs.append("[... truncated]")
                            break
            return "\n".join(paragraphs)[:OFFICE_TEXT_CAP]

        if fmt == "xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(str(file_path), read_only=True, data_only=True)
            lines = []
            row_count = 0
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                lines.append(f"[Sheet: {sheet_name}]")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(cells):
                        lines.append("\t".join(cells))
                        row_count += 1
                        if row_count >= OFFICE_MAX_ROWS:
                            lines.append("[... truncated]")
                            break
                if row_count >= OFFICE_MAX_ROWS:
                    break
            wb.close()
            return "\n".join(lines)[:OFFICE_TEXT_CAP]

    except Exception as e:
        log.warning(f"Office text extraction failed for {file_path.name}: {e}")
    return None


async def describe_media(file_path: str) -> dict:
    """Generate a text description of any media file using Gemini Flash.

    Supports: images, PDFs, video, audio, text, CSV, and office docs (docx/xlsx).
    Returns dict with 'description' on success or 'error' on failure.
    Cost: ~0.001-0.01 USD per call depending on media type and size.

    Size thresholds (configurable via env vars):
    - < INLINE_SIZE_LIMIT (default 20MB): sent inline via Part.from_bytes
    - 20-100MB: uploaded via Gemini File API (auto-deleted after 48h)
    - > MAX_DESCRIBE_SIZE (default 100MB): skipped (cost/time protection)
    """
    start_time = time.time()

    src = Path(file_path)
    if not src.exists():
        return {"error": f"File not found: {file_path}"}

    suffix = src.suffix.lower()

    # Check if we support this file type at all (before requiring API key)
    is_native = suffix in MIME_MAP
    office_fmt = OFFICE_EXTRACTORS.get(suffix)
    if not is_native and not office_fmt:
        return {"error": f"Unsupported file type for description: {suffix}"}

    # Check file size (before requiring API key — fail fast on bad input)
    file_size = src.stat().st_size
    if file_size > MAX_DESCRIBE_SIZE:
        return {"error": f"File too large for description: {file_size / 1024 / 1024:.1f}MB (max {MEDIA_DESCRIBE_MAX_SIZE_MB}MB)"}
    if file_size == 0:
        return {"error": "File is empty"}

    if not GEMINI_API_KEY:
        return {"error": "GEMINI_API_KEY not configured"}

    try:
        from google import genai
        from google.genai import types
    except ImportError as e:
        return {"error": f"Missing package: {e}"}

    category = _get_media_category(suffix)
    prompt = _get_describe_prompt(category, suffix)

    try:
        client = genai.Client(api_key=GEMINI_API_KEY)

        # --- Office docs: extract text, send as text/plain ---
        if office_fmt:
            text_content = _extract_office_text(src, office_fmt)
            if not text_content:
                return {"error": f"Could not extract text from {suffix} file"}
            part = types.Part.from_bytes(
                data=text_content.encode("utf-8"),
                mime_type="text/plain",
            )
        # --- Gemini-native formats: send as bytes or upload ---
        else:
            mime_type = MIME_MAP[suffix]

            if file_size <= INLINE_SIZE_LIMIT:
                # Small file: read into memory and send inline
                raw = src.read_bytes()
                part = types.Part.from_bytes(data=raw, mime_type=mime_type)
            else:
                # Large file: use File API upload (no memory spike — streams to API)
                log.info(f"Large file ({file_size / 1024 / 1024:.1f}MB), using File API upload: {src.name}")
                uploaded = client.files.upload(file=str(src))
                part = uploaded

        # Wrap Gemini call in timeout to prevent indefinite hangs
        response = await asyncio.wait_for(
            asyncio.get_running_loop().run_in_executor(None, lambda: client.models.generate_content(
                model=MODEL_FLASH_TEXT,
                contents=[prompt, part],
            )),
            timeout=MEDIA_DESCRIBE_TIMEOUT_S,
        )

        if response.text:
            # PDFs get full text extraction (20K chars like office docs); others get 1K summary
            desc_cap = OFFICE_TEXT_CAP if suffix == ".pdf" else 1000
            description = response.text.strip()[:desc_cap]
            log.info(f"Media described: {src.name} (type={category}, size={file_size / 1024:.0f}KB, elapsed={time.time()-start_time:.1f}s) -> {len(description)} chars")
            return {"description": description}
        else:
            return {"error": "Gemini returned empty response"}

    except asyncio.TimeoutError:
        log.warning(f"describe_media timeout ({MEDIA_DESCRIBE_TIMEOUT_S}s) for {src.name} ({file_size / 1024 / 1024:.1f}MB)")
        return {"error": f"Gemini API timed out after {MEDIA_DESCRIBE_TIMEOUT_S}s"}
    except Exception as e:
        log.error(f"describe_media error ({src.name}): {e}")
        return {"error": f"Gemini vision error: {str(e)[:300]}"}


TRANSCRIBE_PROMPT = (
    "Transcribe this voice message verbatim. Rules:\n"
    "1. Write the EXACT words spoken, preserving the original language.\n"
    "2. If the speech is NOT in English, add an English translation after "
    "the transcript, prefixed with 'Translation: '.\n"
    "3. If the speech IS in English, just provide the transcript.\n"
    "4. Do NOT add descriptions, timestamps, or commentary.\n"
    "5. If you cannot understand parts, mark them as [inaudible]."
)


async def transcribe_voice(file_path: str) -> dict:
    """Transcribe a voice message using Gemini Flash.

    Synchronous (blocking) — called before agent processing so the agent
    sees the transcript in the prompt. Uses a dedicated prompt optimized
    for verbatim transcription (not search-indexing descriptions).

    No semaphore — this is on the critical path (user is waiting).
    Timeout: VOICE_TRANSCRIBE_TIMEOUT_S (default 60s).

    Returns dict with 'transcript' on success or 'error' on failure.
    """
    start_time = time.time()

    if not GEMINI_API_KEY:
        return {"error": "GEMINI_API_KEY not configured"}

    src = Path(file_path)
    if not src.exists():
        return {"error": f"File not found: {file_path}"}

    suffix = src.suffix.lower()
    mime_type = MIME_MAP.get(suffix)
    if not mime_type:
        return {"error": f"Unsupported audio format: {suffix}"}

    file_size = src.stat().st_size
    if file_size == 0:
        return {"error": "File is empty"}

    try:
        from google import genai
        from google.genai import types
    except ImportError as e:
        return {"error": f"Missing package: {e}"}

    try:
        client = genai.Client(api_key=GEMINI_API_KEY)

        if file_size <= INLINE_SIZE_LIMIT:
            raw = src.read_bytes()
            part = types.Part.from_bytes(data=raw, mime_type=mime_type)
        else:
            log.info(f"Large voice file ({file_size / 1024 / 1024:.1f}MB), using File API upload: {src.name}")
            uploaded = client.files.upload(file=str(src))
            part = uploaded

        response = await asyncio.wait_for(
            asyncio.get_running_loop().run_in_executor(None, lambda: client.models.generate_content(
                model=MODEL_FLASH_TEXT,
                contents=[TRANSCRIBE_PROMPT, part],
            )),
            timeout=VOICE_TRANSCRIBE_TIMEOUT_S,
        )

        if response.text:
            transcript = response.text.strip()
            if len(transcript) > VOICE_TRANSCRIBE_MAX_CHARS:
                transcript = transcript[:VOICE_TRANSCRIBE_MAX_CHARS] + "\n[... truncated]"
            elapsed = time.time() - start_time
            log.info(f"Voice transcribed: {src.name} (size={file_size / 1024:.0f}KB, elapsed={elapsed:.1f}s) -> {len(transcript)} chars")
            return {"transcript": transcript}
        else:
            return {"error": "Gemini returned empty response"}

    except asyncio.TimeoutError:
        elapsed = time.time() - start_time
        log.warning(f"transcribe_voice timeout ({VOICE_TRANSCRIBE_TIMEOUT_S}s) for {src.name} ({file_size / 1024:.0f}KB, elapsed={elapsed:.1f}s)")
        return {"error": f"Transcription timed out after {VOICE_TRANSCRIBE_TIMEOUT_S}s"}
    except Exception as e:
        log.error(f"transcribe_voice error ({src.name}): {e}")
        return {"error": f"Transcription error: {str(e)[:300]}"}
