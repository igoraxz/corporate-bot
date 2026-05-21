# SPDX-License-Identifier: MIT
# Copyright (c) 2026 Corporate Bot Platform
"""Document parser — multi-format content extraction for knowledge domains.

Extracts plain text from various document formats for chunking + embedding:
- Markdown (.md): pass-through (already text)
- PDF (.pdf): Gemini multimodal extraction (fallback: pypdf)
- DOCX (.docx): Gemini multimodal extraction (fallback: python-docx)
- XLSX (.xlsx): Gemini multimodal extraction
- PPTX (.pptx): Gemini multimodal extraction
- Images (.png, .jpg, etc.): Gemini vision description
- Plain text (.txt, .rst, .rtf): pass-through
- Code files (.py, .js, .ts, .go, etc.): pass-through with language hint
- YAML/JSON/TOML (.yaml, .yml, .json, .toml): pass-through
- HTML/XML (.html, .xml): pass-through
- CSV (.csv): pass-through
"""

import logging
from pathlib import Path

log = logging.getLogger(__name__)


# Extension → parser mapping
_PARSERS: dict[str, str] = {
    # Text/docs (pass-through — no LLM needed)
    ".md": "text", ".txt": "text", ".rst": "text", ".rtf": "text",
    ".csv": "text", ".html": "text", ".xml": "text",
    # Structured docs (Gemini multimodal extraction)
    ".pdf": "gemini", ".docx": "gemini", ".xlsx": "gemini",
    ".pptx": "gemini",
    # Images (Gemini vision)
    ".png": "gemini", ".jpg": "gemini", ".jpeg": "gemini",
    ".gif": "gemini", ".webp": "gemini", ".svg": "gemini",
    # Code (pass-through with fence)
    ".py": "code", ".js": "code", ".ts": "code", ".go": "code",
    ".rs": "code", ".java": "code", ".rb": "code", ".sh": "code",
    ".bash": "code", ".css": "code", ".sql": "code",
    # Config (pass-through)
    ".yaml": "text", ".yml": "text", ".json": "text", ".toml": "text",
    ".env": "text", ".ini": "text", ".cfg": "text",
}


def get_parser_type(file_path: str) -> str | None:
    """Get parser type for a file extension. None if unsupported."""
    ext = Path(file_path).suffix.lower()
    return _PARSERS.get(ext)


def is_supported(file_path: str) -> bool:
    """Check if a file format is supported for parsing."""
    return get_parser_type(file_path) is not None


def parse_document(file_path: str, content: bytes | str) -> str:
    """Parse a document into plain text for chunking.

    Args:
        file_path: Original filename (used for extension detection)
        content: File content (bytes for binary formats, str for text)

    Returns: Extracted text content.
    Raises: ValueError if format unsupported.
    """
    parser_type = get_parser_type(file_path)
    if not parser_type:
        raise ValueError(f"Unsupported format: {Path(file_path).suffix}")

    if parser_type == "text":
        return content if isinstance(content, str) else content.decode("utf-8", errors="replace")

    if parser_type == "code":
        text = content if isinstance(content, str) else content.decode("utf-8", errors="replace")
        ext = Path(file_path).suffix.lstrip(".")
        return f"```{ext}\n{text}\n```"

    if parser_type == "gemini":
        # Gemini parsing is async — sync callers use asyncio.run,
        # async callers should use parse_document_async() instead
        import asyncio
        if isinstance(content, str):
            content = content.encode("utf-8")
        try:
            asyncio.get_running_loop()
            # Already in async context — caller must use parse_document_async()
            raise RuntimeError("Use parse_document_async() for Gemini-powered formats in async context")
        except RuntimeError as e:
            if "parse_document_async" in str(e):
                raise
            # No running loop — safe to use asyncio.run
            return asyncio.run(_parse_with_gemini(file_path, content))

    if parser_type == "pdf":
        return _parse_pdf(content)

    if parser_type == "docx":
        return _parse_docx(content)

    return content if isinstance(content, str) else content.decode("utf-8", errors="replace")


def _parse_pdf(content: bytes) -> str:
    """Extract text from PDF using pypdf."""
    try:
        from pypdf import PdfReader
        import io

        if isinstance(content, str):
            content = content.encode("utf-8")

        reader = PdfReader(io.BytesIO(content))
        pages: list[str] = []

        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"[Page {i}]\n{text.strip()}")

        if not pages:
            log.warning("PDF extraction yielded no text (scanned PDF?)")
            return "[PDF with no extractable text — may be scanned/image-based]"

        return "\n\n".join(pages)

    except ImportError:
        log.error("pypdf not installed — cannot parse PDF. pip install pypdf")
        return "[PDF parsing unavailable — install pypdf]"
    except Exception as e:
        log.error("PDF parsing failed: %s", e)
        return f"[PDF parsing error: {e}]"


def _parse_docx(content: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        import io

        if isinstance(content, str):
            content = content.encode("utf-8")

        doc = Document(io.BytesIO(content))
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Convert heading styles to markdown
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                try:
                    hashes = "#" * int(level)
                    parts.append(f"{hashes} {text}")
                except ValueError:
                    parts.append(f"# {text}")
            else:
                parts.append(text)

        # Extract tables
        for table in doc.tables:
            table_rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(cells))
            if table_rows:
                # Add markdown table header separator
                if len(table_rows) > 1:
                    header = table_rows[0]
                    separator = " | ".join("---" for _ in table.rows[0].cells)
                    parts.append(f"\n{header}\n{separator}")
                    parts.extend(table_rows[1:])
                else:
                    parts.extend(table_rows)

        return "\n\n".join(parts) if parts else "[DOCX with no extractable content]"

    except ImportError:
        log.error("python-docx not installed — cannot parse DOCX. pip install python-docx")
        return "[DOCX parsing unavailable — install python-docx]"
    except Exception as e:
        log.error("DOCX parsing failed: %s", e)
        return f"[DOCX parsing error: {e}]"


# ============================================================
# Gemini multimodal extraction
# ============================================================

# MIME type mapping for Gemini multimodal input
_GEMINI_MIME_MAP: dict[str, str] = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
    ".gif": "image/gif", ".webp": "image/webp", ".svg": "image/svg+xml",
}

_GEMINI_EXTRACT_PROMPT = (
    "Extract ALL text content from this document. Preserve structure: "
    "headings as markdown (# H1, ## H2), lists as bullet points, "
    "tables as markdown tables, code as fenced blocks. "
    "For images: describe the visual content in detail. "
    "For spreadsheets: extract all data preserving column/row structure. "
    "For presentations: extract slide titles and content in order. "
    "Output clean, structured markdown text."
)


async def _parse_with_gemini(file_path: str, content: bytes) -> str:
    """Extract text from any document/image using Gemini Flash multimodal.

    Gemini natively handles: PDF, DOCX, XLSX, PPTX, images (PNG, JPG, etc.)
    Returns structured text extraction with section headings preserved.
    Falls back to legacy parsers if Gemini is unavailable.
    """
    import os
    from config import GEMINI_API_KEY as gemini_api_key
    if not gemini_api_key:
        return _fallback_parse(file_path, content)

    try:
        from google import genai
        from google.genai import types
        import asyncio

        client = genai.Client(api_key=gemini_api_key)

        ext = Path(file_path).suffix.lower()
        mime_type = _GEMINI_MIME_MAP.get(ext, "application/octet-stream")

        model_name = os.environ.get("GEMINI_MODEL_DESCRIBE", "gemini-2.5-flash-lite")

        response = await asyncio.to_thread(
            client.models.generate_content,
            model=model_name,
            contents=[
                types.Content(parts=[
                    types.Part(text=_GEMINI_EXTRACT_PROMPT),
                    types.Part(inline_data=types.Blob(mime_type=mime_type, data=content)),
                ]),
            ],
        )

        if response and response.text:
            return response.text

        log.warning("Gemini returned empty response for %s", file_path)
        return _fallback_parse(file_path, content)

    except Exception as e:
        log.warning("Gemini parsing failed for %s: %s — falling back to legacy parser", file_path, e)
        return _fallback_parse(file_path, content)


def _fallback_parse(file_path: str, content: bytes | str) -> str:
    """Fallback to legacy parsers when Gemini is unavailable."""
    ext = Path(file_path).suffix.lower()
    raw = content if isinstance(content, bytes) else content.encode("utf-8")
    if ext == ".pdf":
        return _parse_pdf(raw)
    elif ext == ".docx":
        return _parse_docx(raw)
    elif ext in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg"):
        return f"[Image file: {file_path} — Gemini unavailable for description]"
    elif ext in (".xlsx", ".pptx"):
        return f"[{ext.lstrip('.').upper()} file: {file_path} — Gemini unavailable for extraction]"
    else:
        return content if isinstance(content, str) else content.decode("utf-8", errors="replace")


# ============================================================
# Async document parsing
# ============================================================

async def parse_document_async(file_path: str, content: bytes | str) -> str:
    """Async version of parse_document. Required for Gemini-powered formats.

    Args:
        file_path: Original filename (used for extension detection)
        content: File content (bytes for binary formats, str for text)

    Returns: Extracted text content.
    Raises: ValueError if format unsupported.
    """
    parser_type = get_parser_type(file_path)
    if not parser_type:
        raise ValueError(f"Unsupported format: {Path(file_path).suffix}")

    if parser_type == "gemini":
        if isinstance(content, str):
            content = content.encode("utf-8")
        return await _parse_with_gemini(file_path, content)

    # Non-Gemini formats: delegate to sync parser
    return parse_document(file_path, content)


async def parse_file_content_async(file_path: str, raw_content: bytes | str) -> tuple[str, str]:
    """Async version of parse_file_content. Uses Gemini for supported formats.

    Returns: (parsed_text, chunking_strategy)
    """
    parser_type = get_parser_type(file_path)

    if parser_type == "gemini":
        text = await parse_document_async(file_path, raw_content)
        ext = Path(file_path).suffix.lower()
        # Choose chunking strategy based on content type
        if ext in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg"):
            return text, "recursive"  # Image descriptions are flat text
        elif ext in (".xlsx", ".csv"):
            return text, "recursive"  # Tabular data
        elif ext == ".pptx":
            return text, "markdown_headers"  # Slides have headings
        else:
            return text, "markdown_headers"  # PDF, DOCX
    elif parser_type == "code":
        return parse_document(file_path, raw_content), "code"

    # Non-Gemini: use sync parser
    return parse_file_content(file_path, raw_content)


# ============================================================
# Batch parsing helper
# ============================================================

def parse_file_content(file_path: str, raw_content: bytes | str) -> tuple[str, str]:
    """Parse a file and determine optimal chunking strategy.

    Returns: (parsed_text, chunking_strategy)
    """
    parser_type = get_parser_type(file_path)

    if parser_type == "code":
        return parse_document(file_path, raw_content), "code"
    elif parser_type == "pdf":
        return parse_document(file_path, raw_content), "recursive"
    elif parser_type == "docx":
        return parse_document(file_path, raw_content), "markdown_headers"
    else:
        return parse_document(file_path, raw_content), "markdown_headers"
